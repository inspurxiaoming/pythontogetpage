import urllib

import requests
from bs4 import BeautifulSoup
import os
import docx
from docx import Document
from docx.shared import Inches
from urllib import request
import pypandoc

from url_de import Urlchuli

urllist = [
    # "https://docs.ucloud.cn/security/uws_robot/operation/tasks"
]
From_file = open("../image_url/img_url.txt",'r', encoding='utf-8')
for each_line in From_file:
    if each_line in urllist:
        pass
    else:
        urllist.append(each_line)
conut = 0;
print(len(urllist))
str_count = str(conut)
def getdocName(url_forname,count_forname):
    fp = request.urlopen(url_forname)
    all_html = fp.read()
    soup = BeautifulSoup(all_html, 'html.parser')
    print(soup.find_all("title")[0])
    title_name = soup.find_all("title")
    if len(title_name) > 0 :
        realName = str(title_name[0].get_text())
        if len(soup.find_all("img")) > 2:
            return realName.replace("/","_") + str(count_forname) + "_img.docx"
        else:
            return realName.replace("/","_") + str(count_forname) + ".docx"
    else:
        return str(count_forname) + ".docx"
final_img_name_list = []
for url in urllist:
    docName = getdocName(url,conut)
    html = requests.get(url).content
    soup = BeautifulSoup(html,'html.parser')
    wen = soup.find('div',{"class":"docs-main-content"}).text
    img_list = soup.find_all('img')
    # url = 'https://docs.ucloud.cn/_media/security/uads/程永明.png'
    # de = url.encode('utf-8').decode('utf-8').decode('utf-8')
    for img_url in img_list:
        # print(urllib.parse.unquote(img_url['src'],'utf-8'))
        tu = 'https://docs.ucloud.cn' + urllib.parse.unquote(img_url['src'],'utf-8')
        img_name = tu.split('/')[-1].split('.')[0]
        if 'png' in tu :
            final_img_name = img_name+'.png'
            final_img_name_list.append(final_img_name)
        else:
            final_img_name = img_name +'.jpg'
            final_img_name_list.append(final_img_name)
        #
        # #保存图片至本地
        if 'https://static.ucloud.cn/e69a4f822d7bfeb1cbc4104ab2d1fb3a.jpg' in tu :
            pass
        else:
            with open(final_img_name,'wb')as f:
                response = requests.get(tu).content
                f.write(response)
                f.close()
    #
    document = Document()
    document.add_paragraph(wen)#向文档里添加文字
    # try:
    #     output = pypandoc.convert_file('test.html', 'docx', outputfile="../img_doc/" + docName)
    #     assert output == ""
    # except:
    #     pass

    for image_name in final_img_name_list:
        print(image_name)
        try:
            document.add_picture(image_name)#向文档里添加图片
        except:
            pass
    document.add_paragraph(url)
    document.save('../img_doc/' +docName)#保存文档
    for image_name in final_img_name_list:
        try:
            os.remove(image_name)  # 删除保存在本地的图片
        except:
            pass
    conut +=1
