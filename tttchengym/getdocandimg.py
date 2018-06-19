import requests
from bs4 import BeautifulSoup
import os
import docx
from docx import Document
from docx.shared import Inches

urllist = [
    # "https://docs.ucloud.cn/security/uads/index",
    # "https://docs.ucloud.cn/security/uewaf",
    # "https://docs.ucloud.cn/security/uhas/index",
    # "https://docs.ucloud.cn/security/udas",
    # "https://docs.ucloud.cn/security/uws_robot",
    "https://docs.ucloud.cn/security/uws_robot/operation/tasks"
]
for url in urllist:
    docName = "../a/" + url.split("/")[-1] + ".doc"
    html = requests.get(url).content
    soup = BeautifulSoup(html,'html.parser')
    wen = soup.find('div',{"class":"docs-main-content"}).text
    img = str(soup.find('img')["src"])
    tu = 'https://docs.ucloud.cn' + img
    img_name = img.split('/')[-1]

    #保存图片至本地
    with open(img_name,'wb')as f:
        response = requests.get(tu).content
        f.write(response)
        f.close()

    document = Document()
    document.add_paragraph(wen)#向文档里添加文字
    document.add_picture(img_name)#向文档里添加图片
    document.save(docName)#保存文档
    os.remove(img_name)#删除保存在本地的图片