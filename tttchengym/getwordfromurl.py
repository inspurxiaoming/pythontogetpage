# -*- coding: utf-8 -*-
# import getAlla
from urllib import request
import pypandoc
import docx
from bs4 import BeautifulSoup #导入模块

from getdoc import MYHTMLParser
from tttchengym.geturl import HTMLClient

urllist = [
    # "https://docs.ucloud.cn/security/uads/index",
    # "https://docs.ucloud.cn/security/uewaf",
    # "https://docs.ucloud.cn/security/uhas/index",
    # "https://docs.ucloud.cn/security/udas",
    # "https://docs.ucloud.cn/security/uws_robot",
    # "https://docs.ucloud.cn/security/uads/concepts/overview"
]
From_file = open("allurl.txt")
for each_line in From_file:
    if each_line in urllist:
        pass
    else:
        urllist.append(each_line)
conut = 0;
print(len(urllist))
str_count = str(conut)
def getdocName(url_forname,count_forname):
    fp = request.urlopen(url_forname)
    all_html = fp.read()
    soup = BeautifulSoup(all_html, 'html.parser')
    print(soup.find_all("title")[0])
    title_name = soup.find_all("title")
    if len(title_name) > 0 :
        realName = str(title_name[0].get_text())
        if len(soup.find_all("img")) > 2:
            image_p = open("../image_url/img_url.txt", "a+", encoding='utf-8')
            print(url)
            print(docName)
            image_p.writelines(url)
            return realName.replace("/","_") + "_img.docx"
        else:
            return realName.replace("/","_") + ".docx"
    else:
        return str(count_forname) + ".docx"

# url = "https://docs.ucloud.cn/security/uads/index" #网页地址
for url in urllist:
    docName = getdocName(url,conut)

    wp = request.urlopen(url) #打开连接
    wp = request.urlopen(url) #打开连接
    content = wp.read() #获取页面内容
    fp = open("test.html","w+b") #打开一个文本文件
    fp.write(content) #写入数据
    fp.close() #关闭文件
    try :
        output = pypandoc.convert_file('test.html', 'docx',outputfile = "../a/"+docName)
        assert output == ""
        doc = docx.Document("../a/"+docName)
        doc.add_paragraph(url)
        doc.save("../a/"+docName)
    except :
        errot_p = open("../error/err_url.txt", "a+", encoding='utf-8')
        print(url)
        print(docName)
        errot_p.writelines(url)
    conut+=1
    str_count = str(conut)

# self =''
# MYHTMLParser.__init__(self,docx.Document("d://a.docx"))
# print(self)